"""Estilos e helpers de formatação para os documentos DOCX de reunião
(Pauta e Ata): página A4, uma única página, sem cor/elementos decorativos
— texto preto sobre fundo branco, fonte e espaçamento grandes o
suficiente para preencher a página.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PRETO = RGBColor(0x00, 0x00, 0x00)

FONTE = 'Calibri'
TAM_TITULO = Pt(20)
TAM_SUBTITULO = Pt(13)
TAM_SECAO = Pt(14)
TAM_CORPO = Pt(12)


def novo_documento():
    doc = Document()
    secao = doc.sections[0]
    secao.page_width = Cm(21.0)
    secao.page_height = Cm(29.7)
    secao.top_margin = Cm(1.6)
    secao.bottom_margin = Cm(1.6)
    secao.left_margin = Cm(2.2)
    secao.right_margin = Cm(2.2)

    normal = doc.styles['Normal']
    normal.font.name = FONTE
    normal.font.size = TAM_CORPO
    normal.font.color.rgb = PRETO
    return doc


def titulo_documento(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(texto)
    run.font.name = FONTE
    run.font.size = TAM_TITULO
    run.font.bold = True
    run.font.color.rgb = PRETO
    return p


def subtitulo_documento(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(texto)
    run.font.name = FONTE
    run.font.size = TAM_SUBTITULO
    run.font.color.rgb = PRETO
    return p


def titulo_secao(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(texto)
    run.font.name = FONTE
    run.font.size = TAM_SECAO
    run.font.bold = True
    run.font.color.rgb = PRETO
    return p


def paragrafo(doc, texto, negrito=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(texto)
    run.font.name = FONTE
    run.font.size = TAM_CORPO
    run.font.bold = negrito
    run.font.color.rgb = PRETO
    return p


def bullet(doc, texto):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(texto)
    run.font.name = FONTE
    run.font.size = TAM_CORPO
    run.font.color.rgb = PRETO
    return p


def checklist_item(doc, texto):
    return paragrafo(doc, f'☐ {texto}')


def conteudo_secao(doc, conteudo):
    """Renderiza o conteúdo de uma seção: string vira um parágrafo,
    lista de strings vira bullets."""
    if conteudo is None:
        return
    if isinstance(conteudo, (list, tuple)):
        for item in conteudo:
            bullet(doc, str(item))
    else:
        paragrafo(doc, str(conteudo))


def lista_participantes(doc, participantes):
    """participantes: lista de strings (Pauta) ou de dicts
    {"nome": ..., "presente": bool} (Ata, com marcação ✓/✗)."""
    if not participantes:
        return
    titulo_secao(doc, 'Participantes')
    for participante in participantes:
        if isinstance(participante, dict):
            marca = '✓' if participante.get('presente', True) else '✗'
            texto = f"{marca} {participante['nome']}"
        else:
            texto = str(participante)
        bullet(doc, texto)


def sombrear_celula_preta(celula):
    """Aplica fundo preto a uma célula de tabela (usado no cabeçalho)."""
    tc_pr = celula._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '000000')
    tc_pr.append(shd)


def texto_celula_branco_negrito(celula, texto):
    celula.text = ''
    run = celula.paragraphs[0].add_run(texto)
    run.font.name = FONTE
    run.font.size = TAM_CORPO
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
